#!/usr/bin/env python3
"""Generate docs/aidaegis_ac_remote_user_manual.docx — commercial Chinese user guide."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "aidaegis_ac_remote_user_manual.docx"
IMG = ROOT / "img" / "manual"

C_NAVY = RGBColor(0x18, 0x30, 0x48)
C_NAVY2 = RGBColor(0x24, 0x40, 0x60)
C_TEXT = RGBColor(0x23, 0x26, 0x2A)
C_MUTED = RGBColor(0x6E, 0x74, 0x7C)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def img(name: str) -> Path:
    p = IMG / name
    if not p.is_file():
        raise SystemExit(f"Missing image: {p}")
    return p


def set_run_font(run, size_pt: float, bold: bool = False, color: RGBColor = C_TEXT) -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.font.name = "Microsoft YaHei"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rFonts.set(qn("w:cs"), "Microsoft YaHei")


def set_paragraph_spacing(p, before: float = 0, after: float = 6, line: float = 1.25) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line


def shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color: str = "D2D8DE") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


class ManualDoc:
    def __init__(self) -> None:
        self.doc = Document()
        self._setup_page()
        self._setup_styles()

    def _setup_page(self) -> None:
        section = self.doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("aidaegis AC Remote  ·  用户说明书")
        set_run_font(run, 8, color=C_MUTED)

    def _setup_styles(self) -> None:
        normal = self.doc.styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal.font.size = Pt(10.5)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    def h1(self, text: str) -> None:
        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, before=16, after=8)
        run = p.add_run(text)
        set_run_font(run, 16, bold=True, color=C_NAVY)

    def h2(self, text: str) -> None:
        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, before=12, after=6)
        run = p.add_run(text)
        set_run_font(run, 12.5, bold=True, color=C_NAVY2)

    def p(self, text: str) -> None:
        para = self.doc.add_paragraph()
        set_paragraph_spacing(para, before=0, after=6)
        run = para.add_run(text)
        set_run_font(run, 10.5, color=C_TEXT)

    def bullets(self, items: list[str]) -> None:
        for item in items:
            para = self.doc.add_paragraph(style="List Bullet")
            set_paragraph_spacing(para, before=0, after=3)
            # Clear default and set our run
            if para.runs:
                para.runs[0].text = item
                set_run_font(para.runs[0], 10.5, color=C_TEXT)
            else:
                run = para.add_run(item)
                set_run_font(run, 10.5, color=C_TEXT)

    def numbered(self, items: list[str]) -> None:
        for item in items:
            para = self.doc.add_paragraph(style="List Number")
            set_paragraph_spacing(para, before=0, after=3)
            if para.runs:
                para.runs[0].text = item
                set_run_font(para.runs[0], 10.5, color=C_TEXT)
            else:
                run = para.add_run(item)
                set_run_font(run, 10.5, color=C_TEXT)

    def note(self, text: str) -> None:
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.rows[0].cells[0]
        shade_cell(cell, "F5F8FC")
        set_cell_borders(cell, "B4C8DC")
        para = cell.paragraphs[0]
        set_paragraph_spacing(para, before=4, after=4)
        run = para.add_run(f"提示：{text}")
        set_run_font(run, 9.5, color=C_TEXT)
        self.doc.add_paragraph()

    def caption(self, text: str) -> None:
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(para, before=2, after=10)
        run = para.add_run(text)
        set_run_font(run, 9, color=C_MUTED)

    def fig(self, path: Path, width_cm: float = 10.5, caption: str = "") -> None:
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(para, before=6, after=2)
        run = para.add_run()
        run.add_picture(str(path), width=Cm(width_cm))
        if caption:
            self.caption(caption)

    def figs_row(
        self,
        paths: list[Path],
        captions: list[str],
        width_cm: float = 7.5,
    ) -> None:
        table = self.doc.add_table(rows=2, cols=len(paths))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (path, cap) in enumerate(zip(paths, captions)):
            cell = table.rows[0].cells[i]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(str(path), width=Cm(width_cm))
            set_cell_borders(cell, "FFFFFF")

            cell2 = table.rows[1].cells[i]
            para2 = cell2.paragraphs[0]
            para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(para2, before=2, after=6)
            run2 = para2.add_run(cap)
            set_run_font(run2, 8.5, color=C_MUTED)
            set_cell_borders(cell2, "FFFFFF")
        self.doc.add_paragraph()

    def table(self, headers: list[str], rows: list[list[str]], col_widths_cm: list[float] | None = None) -> None:
        t = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = True

        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            shade_cell(cell, "183048")
            set_cell_borders(cell)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(para, before=3, after=3)
            run = para.add_run(h)
            set_run_font(run, 9.5, bold=True, color=C_WHITE)

        for r_idx, row in enumerate(rows):
            for c_idx, text in enumerate(row):
                cell = t.rows[r_idx + 1].cells[c_idx]
                if r_idx % 2 == 1:
                    shade_cell(cell, "F8FAFC")
                set_cell_borders(cell)
                para = cell.paragraphs[0]
                set_paragraph_spacing(para, before=2, after=2)
                run = para.add_run(text)
                set_run_font(run, 9.5, color=C_TEXT)

        if col_widths_cm:
            for row in t.rows:
                for i, w in enumerate(col_widths_cm):
                    row.cells[i].width = Cm(w)

        self.doc.add_paragraph()

    def page_break(self) -> None:
        self.doc.add_page_break()

    def cover(self) -> None:
        self.fig(img("product_hero_desk_pdf.jpg"), width_cm=16.5)

        for text, size, bold, color, after in [
            ("aidaegis", 26, True, C_NAVY, 2),
            ("AC Remote", 20, True, C_NAVY, 6),
            ("智能空调控制器  ·  用户说明书", 14, False, C_TEXT, 8),
            ("支持 Apple 家庭 / Google Home / Home Assistant", 10.5, False, C_MUTED, 2),
            ("圆屏触控 · Matter 本地控制 · 氛围光环", 10.5, False, C_MUTED, 16),
            ("请妥善保管本说明书，以便日后查阅。", 9.5, False, C_MUTED, 0),
        ]:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=0, after=after)
            run = p.add_run(text)
            set_run_font(run, size, bold=bold, color=color)

        self.page_break()

    def save(self, path: Path) -> Path:
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))
        return path


def build() -> Path:
    m = ManualDoc()
    core = m.doc.core_properties
    core.title = "aidaegis AC Remote 用户说明书"
    core.author = "aidaegis"
    core.subject = "智能空调控制器用户说明书"
    core.keywords = "aidaegis, AC Remote, Matter, 空调, 用户说明书"

    m.cover()

    # TOC
    m.h1("目录")
    toc = [
        ("01", "产品简介"),
        ("02", "外观与界面总览"),
        ("03", "放置建议"),
        ("04", "首次使用：智能家居配网"),
        ("05", "绑定空调（红外学习）"),
        ("06", "圆屏操作说明"),
        ("07", "手机 App 控制"),
        ("08", "氛围灯光"),
        ("09", "恢复出厂设置"),
        ("10", "常见问题"),
        ("11", "产品规格"),
        ("12", "安全与注意事项"),
    ]
    for num, title in toc:
        p = m.doc.add_paragraph()
        set_paragraph_spacing(p, before=2, after=4)
        run = p.add_run(f"{num}    {title}")
        set_run_font(run, 11.5, color=C_TEXT)
    m.page_break()

    # 1
    m.h1("1. 产品简介")
    m.p(
        "感谢您选择 aidaegis AC Remote。本产品是一款面向家庭使用的智能空调控制器，"
        "通过红外方式控制家中空调，并以 Matter 标准接入 Apple「家庭」、Google Home "
        "或 Home Assistant 等智能家居平台，实现本地控制，无需依赖空调厂商云服务。"
    )
    m.p(
        "设备采用光环造型设计：外圈为氛围灯环，内嵌 1.28 英寸圆形触摸屏。"
        "您可在屏上完成配网、空调绑定、日常控制与灯光设置，也可通过手机 App 远程操作。"
    )
    m.h2("主要功能")
    m.bullets(
        [
            "智能家居配网：圆屏显示动态二维码与数字配对码，一键添加至家庭平台",
            "空调绑定：使用原装遥控器完成红外学习，快速对接您的空调",
            "本地控制：开关、制冷/制热、设定温度、风扇档位；屏上与手机状态双向同步",
            "氛围灯环：多种灯光模式与亮度调节，可在手机端开关与调光",
            "环境感知（选配）：支持室温与湿度显示",
            "中英文界面：出厂默认中文，可一键切换英文",
        ]
    )
    m.fig(
        img("product_studio_front_pdf.jpg"),
        width_cm=9.5,
        caption="图 1-1  产品正面（氛围灯环点亮示意）",
    )
    m.note(
        "本产品通过红外协议控制空调。若您的空调遥控协议不受支持，绑定可能失败。"
        "届时请参考「常见问题」或联系售后支持。"
    )

    # 2
    m.h1("2. 外观与界面总览")
    m.p("请先熟悉产品各部分名称，以便后续阅读操作说明。")
    m.h2("2.1 外观结构")
    m.table(
        ["名称", "说明"],
        [
            ["氛围灯环", "外圈环形灯带，提供环境照明与状态氛围"],
            ["圆形触摸屏", "1.28 英寸触控显示屏，用于配网、学习与日常控制"],
            ["底座", "稳定支撑整机，便于桌面放置"],
            ["电源接口", "位于底座后方，连接附赠电源线供电"],
        ],
        [4.0, 12.5],
    )
    m.figs_row(
        [img("product_studio_front_pdf.jpg"), img("product_side_power_pdf.jpg")],
        ["图 2-1  正面外观", "图 2-2  侧面与电源连接"],
        width_cm=7.2,
    )

    m.h2("2.2 圆屏界面一览")
    m.p("设备根据使用阶段自动切换界面，主要页面如下：")
    m.table(
        ["界面", "何时出现", "您可以做什么"],
        [
            ["智能家居配网", "首次上电或未完成配网时", "扫码或输入配对码，加入家庭平台"],
            ["红外学习", "配网完成后尚未绑定空调时", "点按「开始学习」，用原装遥控绑定"],
            ["空调控制", "绑定完成后（日常使用）", "开关空调、调节温度；左滑进入灯光"],
            ["氛围灯光", "在空调页向左滑动进入", "选择灯光模式、调节亮度"],
        ],
        [3.5, 5.0, 8.0],
    )
    m.figs_row(
        [img("ui_pairing_on_device_pdf.jpg"), img("ui_learn_on_device_pdf.jpg")],
        ["图 2-3  配网页", "图 2-4  红外学习页"],
        width_cm=7.2,
    )
    m.figs_row(
        [img("ui_ac_on_device_pdf.jpg"), img("ui_light_on_device_pdf.jpg")],
        ["图 2-5  空调控制页", "图 2-6  氛围灯光页"],
        width_cm=7.2,
    )

    # 3
    m.h1("3. 放置建议")
    m.p("正确放置有助于红外控制更稳定，也能让灯环氛围效果更舒适。")
    m.numbered(
        [
            "将设备放置在空调附近的桌面或台面，保证前方开阔、无遮挡。",
            "使红外发射方向大致朝向空调室内机的接收窗口（通常位于出风口附近）。",
            "使设备能够接收到您日常使用原装遥控器时的红外信号（便于状态同步）。",
            "避免强阳光直射屏幕与接收区域；远离金属大面积遮挡。",
            "使用稳定电源供电；手机与设备须连接同一家庭网络中的 2.4 GHz Wi-Fi。",
        ]
    )
    m.fig(
        img("product_hero_desk_pdf.jpg"),
        width_cm=15.0,
        caption="图 3-1  建议放置于床头柜、书桌等稳定平面",
    )
    m.note(
        "Matter 配网依赖 2.4 GHz Wi-Fi。若家中路由器为双频合一，请确认手机当前已连接 2.4 GHz 频段。"
    )

    # 4
    m.h1("4. 首次使用：智能家居配网")
    m.p(
        "首次上电后，圆屏将显示「Matter 配网」界面，包含动态二维码与数字配对码。"
        "请使用您常用的智能家居 App 完成添加。"
    )
    m.fig(
        img("ui_pairing_on_device_pdf.jpg"),
        width_cm=8.5,
        caption="图 4-1  配网界面：扫码或输入配对码",
    )
    m.h2("操作步骤")
    m.numbered(
        [
            "为设备接通电源，等待圆屏显示「Matter 配网」页面。",
            "确认手机已连接家庭 2.4 GHz Wi-Fi。",
            "打开 Apple「家庭」、Google Home 或 Home Assistant Companion，选择添加 Matter 配件。",
            "扫描屏上二维码；若不便扫码，可手动输入屏上显示的数字配对码。",
            "若系统提示“未认证 / 未经验证的配件”，按提示继续即可。",
            "配网成功后，手机中将出现本设备（默认名称：AC Remote；品牌：aidaegis）。",
        ]
    )
    m.note(
        "屏上二维码由设备实时生成。请务必扫描当前屏幕显示的二维码，"
        "勿使用说明书或其他渠道中的示例图片。"
    )

    # 5
    m.h1("5. 绑定空调（红外学习）")
    m.p(
        "完成智能家居配网后，若尚未绑定空调，屏幕将进入「红外学习」页。"
        "请准备好空调原装遥控器。"
    )
    m.fig(
        img("ui_learn_on_device_pdf.jpg"),
        width_cm=8.5,
        caption="图 5-1  红外学习界面",
    )
    m.h2("操作步骤")
    m.numbered(
        [
            "在圆屏上点按「开始学习」。此时氛围灯环呈黄色柔和呼吸，表示正在等待遥控信号。",
            "将原装遥控器对准本设备，按下任意键（建议：制冷模式、25℃、任意风速）。",
            "学习成功后，屏幕自动进入空调控制页，即可开始使用。",
            "若多次失败，请确认遥控器电量充足、对准方向正确，并参阅第 10 章「常见问题」。",
        ]
    )
    m.note(
        "学习入口仅在触摸屏「开始学习」。绑定完成后，日常使用原装遥控时，"
        "设备可接收信号并尽量与手机端状态保持同步。"
    )

    # 6
    m.h1("6. 圆屏操作说明")
    m.p(
        "绑定完成后，日常控制主要通过空调控制页完成。右上角「EN / 中文」可切换界面语言。"
    )
    m.h2("6.1 空调控制页")
    m.fig(
        img("ui_ac_on_device_pdf.jpg"),
        width_cm=8.5,
        caption="图 6-1  空调控制页示意",
    )
    m.table(
        ["屏幕元素", "操作说明"],
        [
            ["标题「空调」", "表示当前为空调控制界面"],
            ["温度显示", "显示当前设定温度（整度）"],
            ["「开启 / 关闭」", "点按切换空调电源状态"],
            ["「降温」「升温」", "以 1℃ 为步进调节设定温度"],
            ["「EN」语言按钮", "点按切换中文 / 英文界面"],
            ["「左滑灯光」提示", "向左滑动进入氛围灯光页"],
        ],
        [4.5, 12.0],
    )
    m.p("屏上状态与手机智能家居 App 中的控制状态双向同步，任一侧更改都会反映到另一侧。")

    m.h2("6.2 氛围灯光页")
    m.p("在空调控制页向左滑动，进入氛围灯光页；向右滑动可返回空调页。")
    m.fig(
        img("ui_light_on_device_pdf.jpg"),
        width_cm=8.5,
        caption="图 6-2  氛围灯光页示意",
    )
    m.table(
        ["灯光模式", "效果说明"],
        [
            ["夜间关闭", "关闭灯环，适合睡眠环境"],
            ["手动亮度", "白光常亮，亮度跟随滑条或手机调节"],
            ["温感呼吸", "出厂默认；按室温在绿色至橙色间柔和呼吸"],
            ["纯色", "按当前温感色固定点亮"],
            ["彩虹", "色彩循环变化"],
            ["呼吸白", "白光呼吸效果"],
        ],
        [4.0, 12.5],
    )
    m.p(
        "底部亮度滑条可调节灯环明暗。手机 App 可控制灯环开关与亮度；"
        "彩虹、呼吸等氛围模式仅在本机圆屏选择。"
    )

    m.h2("6.3 语言切换")
    m.p(
        "在任意主要界面，点按屏幕右上角「EN」可切换为英文界面；"
        "在英文界面点按「中文」可切回中文。出厂默认中文。"
    )

    # 7
    m.h1("7. 手机 App 控制")
    m.p(
        "完成配网并绑定空调后，即可在手机智能家居应用中控制本设备。"
        "不同平台界面略有差异，可控制内容如下："
    )
    m.table(
        ["控制项", "可操作内容"],
        [
            ["空调", "开关、制冷/制热、目标温度（整度 ℃）"],
            ["风扇", "低 / 中 / 高风速（部分手机界面会单独显示）"],
            ["湿度", "相对湿度显示（需选配温湿度传感）"],
            ["氛围灯", "开关与亮度调节"],
        ],
        [4.0, 12.5],
    )
    m.note(
        "半度温度、扫风摆叶、除湿/仅通风等高级功能，可能受智能家居标准与空调协议限制，"
        "表现可能与原装遥控器不完全一致。设备会尽量将手机指令映射为可用的红外控制。"
    )

    # 8
    m.h1("8. 氛围灯光")
    m.p(
        "灯环既是产品外观的一部分，也可作为柔和的环境照明。"
        "您可在圆屏灯光页选择模式，或在手机 App 中开关灯环并调节亮度。"
    )
    m.figs_row(
        [img("product_studio_front_pdf.jpg"), img("ui_light_on_device_pdf.jpg")],
        ["图 8-1  灯环点亮效果", "图 8-2  灯光设置界面"],
        width_cm=7.2,
    )
    m.bullets(
        [
            "睡眠场景建议选择「夜间关闭」，避免灯光影响休息。",
            "「温感呼吸」适合日常使用，色调随室内温度氛围变化。",
            "学习空调时灯环会临时变为黄色呼吸，学习结束后自动恢复原模式。",
        ]
    )

    # 9
    m.h1("9. 恢复出厂设置")
    m.p(
        "当需要更换家庭网络、重新绑定其他空调，或设备状态异常时，可恢复出厂设置。"
        "恢复后，智能家居配网信息与空调绑定记录将被清除。"
    )
    m.numbered(
        [
            "找到机身复位键，长按约 5 秒。",
            "松开后，设备将清除配网与空调绑定数据并重启。",
            "重新上电后，请按第 4 章重新配网，并按第 5 章重新绑定空调。",
        ]
    )
    m.note("恢复出厂不会损坏硬件。完成后需重新完成配网与空调绑定，方可继续使用手机控制。")

    # 10
    m.h1("10. 常见问题")
    m.table(
        ["现象", "处理建议"],
        [
            ["手机找不到设备", "确认手机连接 2.4 GHz Wi-Fi；靠近路由器后重试；必要时恢复出厂后重新配网"],
            ["屏上二维码扫不进", "确认仍处于配网页；重新上电刷新二维码；或改用数字配对码"],
            ["学习时无反应", "先点「开始学习」再按遥控；对准设备；检查遥控器电池"],
            ["已学习但空调不动", "调整设备朝向空调接收窗；缩短距离、排除遮挡后重试"],
            ["触摸无反应", "使用手指直接触控；避免过厚贴膜；尝试重新上电"],
            ["氛围灯不亮", "确认未选择「夜间关闭」；检查手机端灯是否关闭；适当提高亮度"],
            ["温度/湿度不准或无显示", "温湿度为选配功能；未安装传感器时仍可正常控制空调"],
            ["想换绑另一台空调", "恢复出厂设置后，重新配网并对新空调执行红外学习"],
        ],
        [4.5, 12.0],
    )

    # 11
    m.h1("11. 产品规格")
    m.table(
        ["项目", "说明"],
        [
            ["产品名称", "aidaegis AC Remote"],
            ["显示", "1.28 英寸圆形触摸屏"],
            ["网络", "2.4 GHz Wi-Fi，Matter 本地智能家居控制"],
            ["控制方式", "红外控制空调；圆屏触控；手机 App"],
            ["氛围照明", "环形氛围灯（开关、亮度与多种模式）"],
            ["选配传感", "温湿度传感（如已配置）"],
            ["界面语言", "中文（默认）/ 英文"],
            ["供电", "底座后方电源接口（请使用合格电源适配器）"],
            ["适用环境", "干燥室内桌面放置"],
        ],
        [4.0, 12.5],
    )

    # 12
    m.h1("12. 安全与注意事项")
    m.bullets(
        [
            "仅在干燥的室内环境使用，避免进水、凝露与高湿环境。",
            "请勿遮挡灯环、屏幕及红外收发区域。",
            "请使用符合当地安全标准的电源适配器与线缆。",
            "儿童进行配网或恢复出厂等操作时，请在成人指导下进行。",
            "请勿自行拆解设备；非专业维修可能导致损坏或安全风险。",
            "本说明书仅描述用户日常操作；功能以设备实际表现为准。",
        ]
    )

    p = m.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=24, after=0)
    run = p.add_run(
        "© aidaegis\n"
        "本说明书内容如有更新，以产品随附或官网最新版本为准。\n"
        "品牌与产品名称：aidaegis AC Remote"
    )
    set_run_font(run, 9, color=C_MUTED)

    return m.save(OUT)


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path} ({path.stat().st_size} bytes)")
